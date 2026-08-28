from pathlib import Path
import os
import tempfile
import zipfile

from docx import Document
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
DOCS = [
    ROOT / "docs" / "黑神话悟空_小西天_一座古寺的第二次生命_教案.docx",
    ROOT / "docs" / "黑神话悟空_小西天主题扩展阅读材料包_教师版.docx",
    ROOT / "docs" / "黑神话悟空_小西天主题扩展阅读菜单_学生版.docx",
]

REPLACEMENTS = {
    "2026-06-22": "2025-12-17",
    "https://www.beijing.gov.cn/fuwu/bmfw/bmzt/gjdjbwg/yw/202606/t20260622_4710101.html":
        "https://www.beijing.gov.cn/renwen/zt/gjdjbwg/yw/202512/t20251217_4347718.html",
    "游客流管理既关乎遗产保护，也关乎游客体验与当地社区收益/压力。":
        "通过信息、引导、分区与行为规范管理游客行为，减少对遗产地和参观体验的负面影响。",
}


def iter_paragraphs(document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs_in_cell(cell)
    for section in document.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                yield paragraph
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from iter_paragraphs_in_cell(cell)


def iter_paragraphs_in_cell(cell):
    for paragraph in cell.paragraphs:
        yield paragraph
    for table in cell.tables:
        for row in table.rows:
            for nested_cell in row.cells:
                yield from iter_paragraphs_in_cell(nested_cell)


def replace_in_paragraph(paragraph):
    original = paragraph.text
    updated = original
    for old, new in REPLACEMENTS.items():
        updated = updated.replace(old, new)
    for replacement in REPLACEMENTS.values():
        if replacement.startswith("http"):
            while replacement + replacement in updated:
                updated = updated.replace(replacement + replacement, replacement)
    if updated == original:
        return 0

    if len(paragraph.runs) == 1:
        paragraph.runs[0].text = updated
    else:
        # Source metadata may be split across runs and hyperlink nodes. Clearing
        # the paragraph avoids leaving a second visible copy inside a hyperlink.
        paragraph.clear()
        paragraph.add_run(updated)
    return 1


def replace_in_package(path):
    byte_replacements = [(old.encode("utf-8"), new.encode("utf-8")) for old, new in REPLACEMENTS.items()]
    changed = 0
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", dir=path.parent) as temporary:
        temporary_path = Path(temporary.name)
    try:
        with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temporary_path, "w", zipfile.ZIP_DEFLATED) as target:
            for info in source.infolist():
                data = source.read(info.filename)
                updated = data
                if info.filename.endswith((".xml", ".rels")):
                    for old, new in byte_replacements:
                        updated = updated.replace(old, new)
                if info.filename.endswith(".xml"):
                    root = etree.fromstring(updated)
                    xml_changed = False
                    for paragraph in root.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"):
                        text_nodes = list(paragraph.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"))
                        combined = "".join(node.text or "" for node in text_nodes)
                        collapsed = combined
                        for replacement in REPLACEMENTS.values():
                            if replacement.startswith("http"):
                                while replacement + replacement in collapsed:
                                    collapsed = collapsed.replace(replacement + replacement, replacement)
                        if collapsed != combined and text_nodes:
                            text_nodes[0].text = collapsed
                            for node in text_nodes[1:]:
                                node.text = ""
                            xml_changed = True
                    if xml_changed:
                        updated = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
                if updated != data:
                    changed += 1
                target.writestr(info, updated)
        os.replace(temporary_path, path)
    finally:
        if temporary_path.exists():
            temporary_path.unlink()
    return changed


def main():
    for path in DOCS:
        document = Document(path)
        changes = sum(replace_in_paragraph(p) for p in iter_paragraphs(document))
        for part in document.part.package.parts:
            for relationship in part.rels.values():
                old_target = relationship.target_ref
                new_target = old_target
                for old, new in REPLACEMENTS.items():
                    new_target = new_target.replace(old, new)
                if new_target != old_target:
                    relationship._target = new_target
                    changes += 1
        document.save(path)
        changes += replace_in_package(path)
        print(f"{path.name}: {changes}处修正")


if __name__ == "__main__":
    main()
