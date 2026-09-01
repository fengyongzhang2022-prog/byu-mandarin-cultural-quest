from copy import deepcopy
from pathlib import Path
from docx import Document

SOURCE = Path(r"D:\BYU中文文化任务网站_第二稿\output\doc\5000美元与5万棵树_GAI课程网站_专业审核版.docx")
TEMP = SOURCE.with_name(SOURCE.stem + "_new.docx")
PROJECT_COPY = Path(r"D:\BYU中文文化任务网站_第二稿\sites-app-connector-20205bf7d4e99a89d7154bb849718324-x20\docs\5000美元与5万棵树_GAI课程网站_专业审核版.docx")


def set_paragraph(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell(cell, text):
    paragraph = cell.paragraphs[0]
    set_paragraph(paragraph, text)
    for extra in cell.paragraphs[1:]:
        set_paragraph(extra, "")


def remove_last_rows(table, keep):
    while len(table.rows) > keep:
        row = table.rows[-1]
        table._tbl.remove(row._tr)


def add_row_like(table, values):
    new_tr = deepcopy(table.rows[-1]._tr)
    table._tbl.append(new_tr)
    row = table.rows[-1]
    for cell, value in zip(row.cells, values):
        set_cell(cell, value)


doc = Document(SOURCE)

paragraph_updates = {
    10: "听前入口改为See–Think–Wonder三段口头观察，记录学习者真实的视觉注意、初步解释与问题；每段均以录音完成，并在最后一页与正式讲述对照。",
    13: "人物称呼依据新华社最新报道统一为“赛考斯”。听力材料与教学微纪录片旁白只使用“赛考斯”；本名与英文名仅保留在来源核对说明中。",
    15: "最终口语任务先录制约20秒的“开头—中间—结尾”讲述思路，再获得一条GAI建议；正式讲述至少60秒，不设上限。",
    19: "先用约20秒说明开头、中间和结尾，获得一条AI建议后，面向美国受众完成至少60秒中文讲述（不设上限）。",
    25: "每段听力播放前提供“听前词汇（可选）”按钮，词表默认折叠；第一幕的“地窖／土房、流沙、捆、灌木、乔木”配简明图解。完整听满两遍后开放原文，学习者可划选1—8个汉字查看词级拼音、英文和例句。词内音节连续书写，不加空格。",
    36: "学生听力材料统一使用新华社称呼“赛考斯”。本名罗纳德·萨科尔斯基（Ronald Sakolsky）仅在专家来源核对与原报道说明中出现；真实新闻若保留早期字幕“赛·考斯基”，播放器旁显示统一姓名说明，并保留原始新闻画面。",
    58: "第2页｜听前观察：SEE–THINK–WONDER",
    61: "听前观察",
    62: "页面题目：先看一会儿。你注意到了什么？",
    63: "三段录音：我看见……／我觉得……／我想知道……",
    64: "任务说明（中文）：只根据照片回答。先描述可见细节，再说初步想法，最后提出一个真正想知道的问题。现在不用猜故事，也没有标准答案。",
    65: "Task (English): Look closely at the image. Record what you see, what you think, and what you wonder.",
    67: "进入条件",
    68: "完成“看—想—问”三段短录音后进入第一幕；最后一页将把三段听前录音与完整讲述并置。",
    75: "默认折叠。教师从本幕内容和题目考察点中选词；“地窖／土房、流沙、捆、灌木、乔木”提供简明图解，词条均含词级拼音、英文和贴合原文的简单例句。",
    84: "1985年，19岁的殷玉珍来到毛乌素沙地。她住的土房半埋在沙里，风沙一夜就能把门堵住。第二年，她卖掉一只羊，换回600棵树苗。可是，一场大风以后，活下来的不到10棵。她没有停。她和丈夫用铁锹、扁担和钢钎，把树苗一捆一捆背进沙地，又慢慢学会先固定流沙，再种灌木和乔木。",
    119: "1999年，来华任教的美国教育工作者赛考斯从电视报道里看见殷玉珍的治沙经历。他给多个机构发电子邮件，最终通过基金会筹集5000美元。殷玉珍只留下一张美元作纪念，其余都用来买树苗。2000年春天，赛考斯来到毛乌素，两人见面并共同种下一棵树。那时，殷玉珍已经治沙大约14年。",
    133: "这份帮助看见并支持了已经持续多年的行动。",
    158: "泛听 1｜这一幕主要讲什么？",
    159: "考察点：概括二十多年中治沙行动、公共支持与人物关系的持续发展。三个选项长度和句式接近。",
    168: "长期治沙、公共支持和一段跨国关系都在时间中继续。",
    185: "学习者提示：三段视频彼此独立，可按任意顺序观看。全部看完后，国情卡开放；每张卡同时说明“事实是什么”和“它怎样帮助理解故事”。",
    186: "任务说明（中文）：任选顺序观看三段视频。全部看完后，翻开至少两张国情卡，阅读事实及其与故事的联系，再完成标题信息补充。",
    187: "Task (English): Watch the three videos in any order. After all three, open at least two China context cards, read each fact and its connection to the story, and complete the title-information task.",
    196: "播放与字幕：字幕默认关闭；学习者可按需开启现有字幕，并通过重点词汇表获得语言支架。前两段以1.12倍速播放并保持原音高；第三段保持新闻原速。",
    201: "三段视频彼此独立，可按任意顺序观看。前两段为依据公开资料制作的教学微纪录片，地图、时间轴、数字和人物生活画面包含教学生成／示意重建，不是历史影像；第三段为2026年真实新闻。国情卡在三段视频后补充生活条件、国家工程与科学治沙三个尺度。",
    202: "前两段教学微纪录片提供可选中文字幕并以1.12倍速、保持原音高播放；第三段保持新闻原速。专家需继续核对字幕与旁白对应、关键术语和新闻材料的可理解度。",
    206: "毛乌素沙地位于中国北方，跨越内蒙古和陕西。1978年，中国启动“三北”防护林体系建设工程。“三北”指东北、华北和西北，工程通过植树、种草和保护植被来防风固沙。但到了1985年，毛乌素沙地深处有一个叫井背塘的地方，那里仍然没有公路和电。19岁的殷玉珍住进一间半埋在沙里的土房，风沙一夜就能把门堵住。第二年，她和丈夫卖掉家里唯一的一只羊，换回600多棵树苗。树苗要背进沙地，水要用桶担。一场大风过后，活下来的不到10棵。这样的失败不只发生一次：有一年，5000多棵树苗被沙尘暴掀翻，挖了一个冬天的水渠也几乎被毁掉。他们一遍遍重新栽种，逐渐学会先固定流沙，再种灌木和乔木。十四年后，两人已经种出3万多亩林地。那年10月，在700公里外的洛阳，美国教师赛考斯从电视英语新闻中第一次知道了殷玉珍。电视里的几分钟，把毛乌素十四年的治沙带到了一个远方观众面前。",
    214: "材料性质：中央广播电视总台真实新闻（凤凰网转载）。本站保留真实人物、现场声音和原报道入口。学生材料统一称“赛考斯”；本名与英文名仅在专家来源核对中保留。学习重点：人物姓名、网络寻人、二十六年后重逢、重返毛乌素与再次共同种树。",
    221: "语速、停连和音色是否自然、清楚；可选字幕是否与声音对应；重点词汇表能否支持美国中高级中文学习者理解。",
    228: "三段视频全部看完；翻开至少两张国情卡并读懂“事实—理解作用”；至少选择两项需要补充的背景。",
    237: "第8页｜把故事讲给一个美国听众",
    240: "开头—中间—结尾三段式讲述框架",
    243: "任务说明（中文）：选择一个美国听众。先用约20秒说明开头、中间和结尾怎样安排，查看一条AI建议，再完成至少60秒的正式讲述。正式讲述不设上限；可以只录声音，也可以自愿录制画面和声音。",
    244: "Task (English): Choose one U.S. audience. Record a brief plan for the beginning, middle, and ending, review one AI suggestion, and then tell the full story for at least 60 seconds. There is no maximum length. You may record audio only or voluntarily record both video and audio.",
    245: "跟读热身（至少5秒）：这五千美元买来了新的树苗，也支持了殷玉珍已经坚持多年的治沙行动。",
    246: "思路录音：约20秒；说清开头、中间和结尾，考虑美国听众需要补充的治沙背景。",
    247: "AI只给一条可以立即使用的口语建议；优先检查三段结构、时间线、国情背景与5000美元的作用。",
    248: "正式讲述：至少60秒，不设上限；至少70个汉字，包含两个时间点，并说明5000美元的作用和治沙行动怎样长期发展。",
    253: "任务说明（中文）：比较听前“我看见／我觉得／我想知道”的三段录音和最后的讲述，找出你补充了哪些时间、证据、背景和受众信息。",
    254: "Task (English): Compare your See–Think–Wonder recordings with your final telling. Identify the time, evidence, background, and audience information you added.",
    268: "不得冒充赛考斯本人，不生成未公开的内心活动、私人经历或虚构原话；本名与英文名只用于来源核对。",
}

for index, text in paragraph_updates.items():
    set_paragraph(doc.paragraphs[index], text)

# 九页总览
t = doc.tables[1]
set_cell(t.cell(2, 1), "听前观察")
set_cell(t.cell(2, 2), "森林图片＋See–Think–Wonder三段录音")
set_cell(t.cell(2, 3), "录下可见细节、初步解释与真实问题")
set_cell(t.cell(8, 2), "受众＋开头／中间／结尾结构＋思路录音")
set_cell(t.cell(8, 3), "约20秒思路；AI建议；至少60秒正式讲述")
set_cell(t.cell(9, 2), "听前看—想—问与最终讲述对照")

# 第2页 See–Think–Wonder
t = doc.tables[4]
for r, values in enumerate([
    ("步骤", "口头任务"),
    ("SEE｜我看见……", "描述照片中能够看见的细节；至少3秒"),
    ("THINK｜我觉得……", "说出照片引发的想法；至少4秒"),
    ("WONDER｜我想知道……", "提出一个真正想知道的问题；至少4秒"),
]):
    set_cell(t.cell(r, 0), values[0]); set_cell(t.cell(r, 1), values[1])

# 第一幕听前词表：补“捆”
t = doc.tables[5]
add_row_like(t, ["捆", "kǔn", "bundle; to tie"])

# 第一幕三道题，与网站一致
for row, values in enumerate([
    ("A", "殷玉珍在第一次失败后调整方法、继续治沙", "✓ 正确"),
    ("B", "一笔海外捐款解决了树苗问题", ""),
    ("C", "当地已有成熟道路和机械", ""),
], 1):
    for col, value in enumerate(values): set_cell(doc.tables[6].cell(row, col), value)

# 第二幕选项缩短、平行
for row, values in enumerate([
    ("A", "一份善意进入长期治沙行动", "✓ 正确"),
    ("B", "一位教师规划当地生态治理", ""),
    ("C", "殷玉珍第一次决定种树", ""),
], 1):
    for col, value in enumerate(values): set_cell(doc.tables[12].cell(row, col), value)
for row, values in enumerate([
    ("A", "钱被全部保存下来作纪念", ""),
    ("B", "钱用于买树苗；见面时她已治沙约14年", "✓ 正确"),
    ("C", "两人见面以后才开始筹款", ""),
], 1):
    for col, value in enumerate(values): set_cell(doc.tables[14].cell(row, col), value)

# 第三幕泛听题与两个精听题
for table_index, rows in {
    18:[("A", "治沙行动在多年中继续发展", "✓ 正确"),("B", "一次旅行改变了两人的计划", ""),("C", "一场采访带来了商业合作", "")],
    19:[("A", "整个毛乌素的全部树木", ""),("B", "那批捐款树苗后来形成的规模", "✓ 正确"),("C", "殷玉珍种下的全部树木", "")],
    20:[("A", "两人重逢并再次共同种树", "✓ 正确"),("B", "两人第一次通过电视认识", ""),("C", "当地第一次修通公路", "")],
}.items():
    table = doc.tables[table_index]
    for row, values in enumerate(rows, 1):
        for col, value in enumerate(values): set_cell(table.cell(row, col), value)

# 国情卡：直接写出卡片内容和理解作用
t = doc.tables[26]
remove_last_rows(t, 4)
for r, values in enumerate([
    ("时代", "国情卡内容"),
    ("1980年代", "事实：井背塘没有公路和电，树苗靠人背，水靠扁担和水桶运。帮助理解：治沙最初也是日常生存问题。"),
    ("1978—2050", "事实：“三北”指东北、华北和西北，工程持续开展防风固沙和植被保护。帮助理解：个人行动发生在中国北方长期生态治理进程中。"),
    ("今天", "事实：根据水资源与土地条件选择乔木、灌木或草，并配合道路、机械和监测。帮助理解：成活来自方法调整和长期维护。"),
]):
    set_cell(t.cell(r, 0), values[0]); set_cell(t.cell(r, 1), values[1])

# 第8页三段式结构，替换六镜头表
t = doc.tables[29]
remove_last_rows(t, 4)
for r, values in enumerate([
    ("部分", "时间", "讲述重点"),
    ("开头", "1985—1986", "生活条件、第一次失败，以及她为什么继续"),
    ("中间", "1999—2000", "电视报道、5000美元和两人共同种树"),
    ("结尾", "二十多年—2026", "长期治沙、时代变化、寻找与重逢"),
]):
    for c, value in enumerate(values): set_cell(t.cell(r, c), value)

t = doc.tables[30]
set_cell(t.cell(1, 0), "听故事以前：SEE／THINK／WONDER三段录音")
set_cell(t.cell(1, 1), "至少60秒：面向美国受众的完整讲述（无上限）")

# 技术门槛表改成当前流程
t = doc.tables[34]
remove_last_rows(t, 4)
for r, values in enumerate([
    ("任务", "时长", "内容门槛"),
    ("标题说明", "至少12秒", "补充一个年代背景和一个故事证据"),
    ("讲述思路", "约20秒", "开头、中间、结尾；考虑美国听众"),
    ("正式讲述", "至少60秒；无上限", "≥70汉字；两个时间点；5000美元的作用与长期治沙"),
]):
    for c, value in enumerate(values): set_cell(t.cell(r, c), value)

# 完整词表补“捆”，更新三北工程例句
t = doc.tables[35]
for row in t.rows:
    if row.cells[0].text.strip() == "三北工程":
        set_cell(row.cells[3], "三北指东北、华北和西北，工程持续开展防风固沙与生态建设。")
if not any(row.cells[0].text.strip() == "捆" for row in t.rows):
    add_row_like(t, ["捆", "kǔn", "bundle; to tie", "她把树苗一捆一捆背进沙地。"])

doc.save(TEMP)

# 重新打开确认结构完整，再原子替换两个稳定文件。
check = Document(TEMP)
assert len(check.paragraphs) == len(doc.paragraphs)
assert check.paragraphs[62].text == "页面题目：先看一会儿。你注意到了什么？"
assert any(row.cells[0].text.strip() == "捆" for row in check.tables[35].rows)
TEMP.replace(SOURCE)
PROJECT_COPY.write_bytes(SOURCE.read_bytes())
print(SOURCE)
print(PROJECT_COPY)
