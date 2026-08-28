from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"

FILES = {
    "dialogue": DOCS / "黑神话悟空_GAI沉浸式文化学习产品_对话与演示方案_精修版.docx",
    "lesson": DOCS / "黑神话悟空_一件不该被抢的衣服_聚焦版方案_精修版.docx",
    "architecture": DOCS / "黑神话悟空_沉浸式体验架构_交接规格_精修版.docx",
    "review": DOCS / "黑神话悟空_袈裟主题教学价值与跨文化支架_教案审议稿.docx",
    "teacher": DOCS / "黑神话悟空_袈裟主题扩展阅读材料包_教师版.docx",
    "student": DOCS / "黑神话悟空_袈裟主题扩展阅读菜单_学生版.docx",
}

INK = "352A23"
PAPER = "F7F0E3"
PAPER_DARK = "E9D9BE"
CINNABAR = "A33F2C"
GOLD = "B58A45"
GOLD_DARK = "765B32"
MOSS = "526153"
ASH = "73685E"
WHITE = "FFF9ED"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def borders(cell, color="C7AD85", size="5"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = tc_borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def run_style(run, size=None, bold=None, color=None, east="等线", italic=None):
    run.font.name = "Aptos"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_document(doc, footer):
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.05)
    section.right_margin = Cm(2.05)
    styles = doc.styles
    for name in ["Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        if name not in styles:
            styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if "Table Grid" not in styles:
        styles.add_style("Table Grid", WD_STYLE_TYPE.TABLE)
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "等线")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.line_spacing = 1.38
    normal.paragraph_format.space_after = Pt(5)
    for name, size, color, before in [
        ("Title", 25, INK, 0),
        ("Subtitle", 12, CINNABAR, 2),
        ("Heading 1", 16.5, CINNABAR, 13),
        ("Heading 2", 12.5, GOLD_DARK, 10),
        ("Heading 3", 11, MOSS, 8),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "思源宋体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(5)
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_style(footer_p.add_run(footer), 8, color=ASH)


def cover(doc, label, title, subtitle, lead):
    band = doc.add_table(rows=1, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = band.cell(0, 0)
    shade(cell, INK)
    borders(cell, INK, "0")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run_style(p.add_run(label.upper()), 9, True, GOLD)
    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_before = Pt(18)
    p.add_run(title)
    p = doc.add_paragraph(style="Subtitle")
    p.add_run(subtitle)
    rule = doc.add_table(rows=1, cols=2)
    rule.autofit = False
    rule.columns[0].width = Cm(3.2)
    rule.columns[1].width = Cm(13.4)
    shade(rule.cell(0, 0), CINNABAR)
    shade(rule.cell(0, 1), GOLD)
    for c in rule.rows[0].cells:
        borders(c, c._tc.get_or_add_tcPr().find(qn("w:shd")).get(qn("w:fill")), "0")
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(13)
    run_style(p.add_run(lead), 11.5, True, INK, "思源宋体")


def callout(doc, title, text, color=PAPER_DARK):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, color)
    borders(cell, GOLD)
    p = cell.paragraphs[0]
    run_style(p.add_run(title + "　"), 10.5, True, CINNABAR)
    run_style(p.add_run(text), 10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = t.rows[0]
    repeat_header(header)
    for i, value in enumerate(headers):
        cell = header.cells[i]
        cell.text = value
        shade(cell, GOLD_DARK)
        borders(cell, GOLD_DARK)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in cell.paragraphs[0].runs:
            run_style(r, 9, True, WHITE)
    for row_i, values in enumerate(rows):
        row = t.add_row()
        for i, value in enumerate(values):
            cell = row.cells[i]
            cell.text = str(value)
            shade(cell, PAPER if row_i % 2 == 0 else "EFE2CE")
            borders(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                for r in p.runs:
                    run_style(r, 9, color=INK)
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(.55)
        p.paragraph_format.first_line_indent = Cm(-.3)
        p.add_run(item)


def build_dialogue():
    doc = Document()
    style_document(doc, "BYU MANDARIN · 袈裟文化任务 · 对话与演示")
    cover(doc, "BLACK MYTH: WUKONG · CULTURAL QUEST 01", "一件衣服，两种欲望", "GAI角色对话与三分钟现场演示", "学生进入火后的观音禅院，以“记事人”的身份追查袈裟在金池长老与黑熊精眼中的意义。")

    doc.add_heading("一、核心矛盾", level=1)
    doc.add_paragraph("袈裟承担修行身份与少欲提醒；《西游记》中的锦襕袈裟又以珍贵材料和文学想象吸引观看、收藏与争夺。金池长老收藏袈裟，仍然想占有唐僧的锦襕袈裟；黑熊精趁火取走袈裟，又把它视为值得展示的宝物。学生围绕同一件衣服追问：价值从哪里来，喜欢怎样转成占有，占有怎样推动伤害。")
    callout(doc, "演示主问题", "一件提醒人放下欲望的衣服，为什么会成为欲望的中心？")

    doc.add_heading("二、学习者与语言表现", level=1)
    table(doc, ["项目", "要求"], [
        ("对象", "美国大学 Intermediate High–Advanced Low 中文学习者"),
        ("词汇", "袈裟、衣服、珍贵、收藏、占有、放下、欲望、价值、身份、后果"),
        ("句法", "本来……后来……；因为……所以……；它们都……可是……；这个例子能说明……"),
        ("口语", "提问、追问、回应角色、比较物件、完成45–60秒文化解释"),
        ("思维", "从物的外观进入人物行动，再形成有关价值与欲望的判断"),
    ], [3.1, 13.5])

    doc.add_heading("三、三分钟演示", level=1)
    table(doc, ["时间", "体验", "画面与声音", "学习者行动"], [
        ("0:00–0:25", "Enter", "焦木案台上的袈裟；低风、远钟、余火", "接受记事任务，选择是否开启环境声与摄像头"),
        ("0:25–0:55", "Discover", "袈裟外观、收藏记录、人物行动三条短证据", "拾取一条证据，说出一个发现或疑问"),
        ("0:55–1:35", "Interact", "金池长老苍老男声；环境声自动压低", "连续语音对话；追问他为何仍想要这件衣服"),
        ("1:35–2:05", "Compare", "黑熊精低沉男声；火光增强", "用证据回应“宝物就该归最强的人”"),
        ("2:05–2:45", "Explain", "镜头回到袈裟；背景完全静音", "完成45–60秒解释"),
        ("2:45–3:00", "Reflect", "系统显示一条口语建议", "补说15–20秒；进入持续开放的资料架"),
    ], [2.3, 2.5, 6.1, 5.7])

    doc.add_heading("四、角色对话", level=1)
    table(doc, ["角色", "开场", "追问方向", "语音"], [
        ("金池长老", "我守了二百多年，也见过许多袈裟。你说，我为什么偏偏忘不了这一件？", "收藏、身份、他人眼光、喜欢与占有", "苍老、缓慢、短句；每轮一句回应和一个问题"),
        ("黑熊精", "既然人人都说它是宝物，强者拿走它，有什么错？", "物的价值、行动后果、拥有资格、比较边界", "低沉、直接、略带不服；愿意根据证据改口"),
    ], [3.0, 5.4, 5.5, 3.5])
    doc.add_paragraph("AI先回应学生刚说的话，再推进一个问题。学生可反复追问；角色每句回应均可重听。页面保留完整的语音轮次、转写、角色回应和时间顺序。")

    doc.add_heading("五、随时提示", level=1)
    table(doc, ["层级", "提示"], [
        ("方向", "看看这件衣服原来的功能、后来获得的价值、想得到它的人和行动后果。"),
        ("追问", "你已经有很多件，为什么还想要这一件？／你喜欢的是衣服，还是别人看你的方式？"),
        ("比较", "牛仔裤／限量运动鞋／棒球帽／你文化中的一件物品。"),
        ("组织", "本来……后来……我认为……／它们都……可是……"),
    ], [3.0, 13.6])

    doc.add_heading("六、预期口语", level=1)
    callout(doc, "教师参照", "袈裟本来可以提醒人少一点欲望。锦襕袈裟看起来很珍贵，金池长老已经收藏了很多件，还是想得到它。牛仔裤也会从工作服变成身份和价格的象征；它的原来功能和袈裟不同。衣服会有意义，人也会把欲望放进衣服里。", "E6D7BF")
    doc.add_paragraph("这段话用于判断任务难度。学生保留自己的用词、证据顺序和文化例子。")

    doc.add_heading("七、研究记录", level=1)
    bullets(doc, [
        "前测录音45–60秒，最长3分钟；学生可回听、重录并保留最终版本。",
        "对话过程保存学生原始语音、自动转写、AI角色回复、提示使用和轮次时间。",
        "Explain保存45–60秒语音或可选视频；学生可回看、重说。",
        "前后文化解释比较关注证据、因果、跨文化比较、限定表达和口语连贯性。",
        "任务完成后记录扩展阅读选择与新增60–90秒口语说明。",
    ])
    doc.save(FILES["dialogue"])


def build_lesson():
    doc = Document()
    style_document(doc, "BYU MANDARIN · 一件不该被抢的衣服 · 聚焦教案")
    cover(doc, "FOCUSED LESSON · INTERMEDIATE HIGH–ADVANCED LOW", "一件不该被抢的衣服", "《黑神话：悟空》袈裟文化任务", "学生观察袈裟、追问人物、比较文化物件，完成一段有证据的中文解释。")

    doc.add_heading("一、主题内容", level=1)
    doc.add_paragraph("本任务围绕《黑神话：悟空》相关叙事中的袈裟展开。袈裟是一件可以观察材料、颜色、缝制方式和使用者身份的衣服；锦襕袈裟在文学叙事中又成为珍贵、收藏和争夺的对象。它同时连接修行提醒、社会眼光、个人欲望和行动后果。")
    callout(doc, "文化理解", "物有功能，也会被人赋予身份、价格和情感。人物怎样看物，决定了他们怎样行动。")

    doc.add_heading("二、学习目标", level=1)
    table(doc, ["维度", "学习者能够"], [
        ("识别", "说出袈裟的外观、用途和故事中的特殊价值"),
        ("关系", "用本来、后来、因为、所以连接物、人物和行动"),
        ("互动", "向金池长老与黑熊精提问，回应角色的新问题"),
        ("比较", "用熟悉的文化物件说明一个相似点和一个差异"),
        ("解释", "完成45–60秒中文说明，并根据一条建议补说"),
    ], [3.1, 13.5])

    doc.add_heading("三、文化事实卡", level=1)
    table(doc, ["证据", "学生可见文本", "来源类型"], [
        ("袈裟外观", "早期佛教传统中的粪扫衣可由弃布清洗、缝成。拼接外观能提示朴素与节制。", "博物馆藏品说明与教师改写"),
        ("锦襕袈裟", "《西游记》把唐僧的锦襕袈裟写成缀有珍宝的特殊衣物。", "文学原典"),
        ("收藏记录", "金池长老已经收藏许多袈裟，看见锦襕袈裟后仍想占有。", "文学叙事与游戏人物线索"),
        ("行动后果", "占有欲推动借看、藏留、纵火与趁火取衣。", "文学情节概括"),
    ], [3.1, 9.2, 4.3])

    doc.add_heading("四、课堂流程", level=1)
    table(doc, ["环节", "时间", "教学输入", "口语产出"], [
        ("前测", "课前", "“这件衣服为什么会被抢？”；录音上限3分钟", "回听并保留一个版本"),
        ("Enter", "30秒", "火后禅院与焦木案台上的袈裟", "接受记事任务"),
        ("Discover", "60–90秒", "四条证据自由拾取；随时进入角色对话", "一个发现＋一个问题"),
        ("Interact", "1–2分钟", "金池长老围绕收藏、身份与欲望追问", "两轮以上连续口语"),
        ("Compare", "1–2分钟", "黑熊精提出“宝物归强者”；跨文化物件菜单", "证据回应＋相似点和差异"),
        ("Explain", "45–60秒", "本来／后来／我认为三枚提示词", "完整文化解释"),
        ("Reflect", "20秒", "GAI给一条建议", "补说一个修改"),
        ("Extend", "课后随时", "原著、文物、游戏改编、新闻四条路径", "60–90秒新增说明"),
    ], [2.7, 2.2, 6.1, 5.5])

    doc.add_heading("五、差异化支架", level=1)
    table(doc, ["需要", "系统提供", "研究记录"], [
        ("寻找关系", "功能／价值／人物／行动四个观察轴", "打开时间与采用信息"),
        ("组织提问", "为什么还想要？／谁让它变得珍贵？／得到以后发生了什么？", "提示后下一轮话语"),
        ("跨文化比较", "牛仔裤、限量运动鞋、棒球帽、学习者自选物件", "选用物件与类比边界"),
        ("连续表达", "本来……后来……；它们都……可是……", "连接词与自我修正"),
    ], [3.0, 8.0, 5.6])
    doc.add_paragraph("Hint入口在每个环节保持可见。提示逐层展开，学生决定使用深度。")

    doc.add_heading("六、评价标准", level=1)
    table(doc, ["维度", "达标", "发展充分"], [
        ("证据", "提到一条具体材料", "整合两条证据并说明关系"),
        ("因果", "连接人物与行动", "解释价值、欲望与后果"),
        ("互动", "回应角色问题", "主动追问并根据回应调整表达"),
        ("比较", "说出相似或差异", "说明相似、差异与类比边界"),
        ("中文", "形成基本连贯的话段", "连接清楚，能够补充和自我修正"),
    ], [3.0, 6.8, 6.8])

    doc.add_heading("七、事实表达", level=1)
    bullets(doc, [
        "“早期佛教传统中的粪扫衣可由弃布清洗、缝成。”",
        "“锦襕袈裟是《西游记》文学叙事中的特殊宝衣。”",
        "“金池长老与黑熊精呈现两种与袈裟相关的占有行动。”",
        "“牛仔裤等物件帮助比较功能、身份和价格；袈裟还带有宗教身份与修行意义。”",
        "“游戏图像提供人物与气氛；文化事实来自标明来源的材料。”",
    ])

    doc.add_heading("八、教学后延伸", level=1)
    doc.add_paragraph("任务页长期保留四条扩展路径。学生按兴趣阅读一条材料，把新证据加入原来的解释。教师可比较即时解释与延后解释，观察文化理解、口语组织和类比边界的发展。")
    doc.save(FILES["lesson"])


def build_architecture():
    doc = Document()
    style_document(doc, "BYU MANDARIN · 袈裟文化任务 · 环境架构")
    cover(doc, "IMPLEMENTATION SPECIFICATION · 2026-08-28", "袈裟文化任务：沉浸式环境架构", "访问、画面、声音、语音、数据与移动端交接规格", "同一件袈裟贯穿进入、发现、角色互动、跨文化比较、口语解释与课后延伸。")

    doc.add_heading("一、正式访问地址", level=1)
    table(doc, ["入口", "地址", "用途"], [
        ("学生端", "https://quest.usmandarincurriculumlab.com/heishenhuawukong.html", "核心任务、语音角色互动、扩展阅读"),
        ("教师端", "https://quest.usmandarincurriculumlab.com/teacher.html", "会话记录、语音回听、人工编码与导出"),
        ("登录", "https://quest.usmandarincurriculumlab.com/login.html", "进入学生端或教师端"),
    ], [2.7, 9.6, 4.3])

    doc.add_heading("二、体验结构", level=1)
    table(doc, ["阶段", "主画面", "操作", "持续状态"], [
        ("Enter", "焦木案台上的袈裟", "输入编号；检查麦克风；选择环境声与摄像头", "建立会话"),
        ("Pre", "袈裟近景", "录音、回听、重录；最长3分钟", "保留最终版音频与转写"),
        ("Discover", "袈裟外观与四条短证据", "自由拾取；随时去问角色；随时返回", "记录证据与问题"),
        ("Interact", "金池长老游戏场景", "点击圆点开始／结束语音；重听角色；连续追问", "保存双方语音与文本顺序"),
        ("Compare", "黑熊精火场游戏场景", "语音协商；打开跨文化物件Hint", "记录改口与比较"),
        ("Explain", "袈裟回到画面中心", "语音或可选摄像头视频；回看、重说", "保存前后版本"),
        ("Extend", "山外见闻资料架", "任务后随时阅读、补录", "材料长期开放"),
    ], [2.7, 4.9, 6.1, 4.0])

    doc.add_heading("三、视觉系统", level=1)
    table(doc, ["项目", "规范"], [
        ("视觉主线", "焦木、残火、旧绢、暗金；袈裟承担每个阶段的视觉锚点"),
        ("主视觉", "生成的袈裟物件图用于Enter、Pre、Discover与Explain；金池长老和黑熊精使用对应游戏截图"),
        ("配色", "炭黑 #0B0D0C；旧纸 #E8D8BC；朱砂 #B54A32；暗金 #B88A42；灰绿 #667166"),
        ("字体", "中文标题使用宋体类衬线；正文使用系统黑体；英文仅作小号环境标签"),
        ("布局", "桌面端左侧任务与主视觉，右侧阶段导航和Hint；移动端单列，主要操作保持拇指可达"),
        ("动效", "烟尘缓移、火光轻微呼吸、阶段淡入；单次转场180–320ms"),
    ], [3.1, 13.5])

    doc.add_heading("四、声音与音乐", level=1)
    table(doc, ["场景", "声音层", "行为"], [
        ("进入", "原创低风、远钟、余火", "用户点击“开启环境声”后播放；浏览器保持静音进入"),
        ("发现", "低风＋纸张拾取声", "打开证据时轻响；音量低于角色语音20dB以上"),
        ("金池长老", "苍老男声", "角色说话前环境声自动压低；结束后缓慢恢复"),
        ("黑熊精", "低沉男声＋远火", "火声保持背景层；角色句子可重听"),
        ("学生说话", "环境声与界面音归零", "录音开始立即静音；录音结束后等待AI回应"),
        ("Explain", "全静音", "保留学生原声；视频录制同样执行"),
    ], [3.0, 6.0, 7.6])
    callout(doc, "音频来源", "环境声由网页实时合成，避免外部音乐文件加载和循环接缝。角色语音优先选择设备中的普通话男声；系统记录实际声音名称。")

    doc.add_heading("五、系统要求", level=1)
    table(doc, ["项目", "最低要求", "推荐"], [
        ("访问", "HTTPS；稳定网络", "下行5 Mbps以上"),
        ("桌面浏览器", "Chrome或Edge近两个主版本", "最新版Chrome或Edge"),
        ("移动端", "iOS Safari 17+；Android Chrome", "Android Chrome最新版"),
        ("权限", "麦克风；摄像头可选", "首次进入时完成权限检查"),
        ("设备", "扬声器与麦克风", "耳机或安静环境，减少回声"),
        ("语音识别", "浏览器普通话识别服务可用", "Chrome/Edge普通话识别"),
        ("本地容量", "至少100MB可用空间", "任务结束后及时导出研究数据"),
    ], [3.0, 6.5, 7.1])

    doc.add_heading("六、语音交互", level=1)
    bullets(doc, [
        "圆点第一次点击开始录音，第二次点击结束并发送；按钮状态、计时和麦克风电平同时变化。",
        "学生语音以WebM保存到IndexedDB；转写进入会话记录；AI使用转写生成角色回应。",
        "学生对话区显示“语音已发送”和角色回复；角色每句均有重听按钮。",
        "前测与Explain提供回听、重录和最终版本确认；角色对话保持连续轮次。",
        "摄像头作为沉浸窗口单独开关；角色语音对话由圆点控制。",
        "识别空白时保留录音状态并引导学生再说一轮；页面保持当前阶段与历史记录。",
    ])

    doc.add_heading("七、Hint", level=1)
    table(doc, ["层级", "内容", "阶段"], [
        ("方向", "功能、价值、人物、行动、后果", "全程"),
        ("追问", "为什么还想要？谁让它珍贵？得到以后发生什么？", "Discover／Interact"),
        ("比较", "牛仔裤、限量运动鞋、棒球帽、自选文化物件", "Compare"),
        ("句框", "本来……后来……；它们都……可是……", "Compare／Explain"),
    ], [2.6, 9.2, 4.8])
    doc.add_paragraph("Hint逐层展开，学生主动选择。系统记录层级、打开时间、停留时长和下一轮口语变化。")

    doc.add_heading("八、数据结构", level=1)
    table(doc, ["数据", "保存内容", "教师端"], [
        ("会话", "参与者编号、开始／结束时间、完成阶段", "状态与完成度"),
        ("前测", "最终录音、时长、转写、重录次数", "回听与转写"),
        ("证据", "打开顺序、停留、学习者发现与问题", "过程分析"),
        ("对话", "角色、轮次、双方文本、学生音频键、时长", "顺序回放与编码"),
        ("Hint", "阶段、层级、时间、下一轮文本", "支架使用分析"),
        ("解释", "初次与补说语音／视频、转写、语言标记", "前后对照"),
        ("延伸", "材料选择、返回时间、新增口语", "延后学习分析"),
    ], [2.8, 9.1, 4.7])

    doc.add_heading("九、移动端", level=1)
    bullets(doc, [
        "视口使用100svh；键盘弹出时主要录音按钮保持可见。",
        "触控目标达到44×44px；底部操作区避开iPhone安全区域。",
        "摄像头窗口保持16:9，学生可折叠；对话历史独立滚动。",
        "背景图采用响应式裁切；袈裟在窄屏中保持可见。",
        "页面切入后台时停止环境声并暂停摄像头预览。",
    ])

    doc.add_heading("十、验收", level=1)
    table(doc, ["检查", "通过标准"], [
        ("地址", "学生端、教师端与登录地址可访问"),
        ("内容", "页面围绕袈裟、金池长老、黑熊精与跨文化物件展开"),
        ("声音", "环境声需用户开启；学生说话时静音；角色句子可重听"),
        ("录音", "桌面与手机均可开始、结束、保存；前测可回听重录"),
        ("AI", "两名角色回应当前学生话语，每轮一个问题"),
        ("导航", "每个阶段可返回；历史与已选证据保留"),
        ("Hint", "全程可打开；逐层显示；使用进入研究记录"),
        ("延伸", "任务完成后资料架持续开放"),
        ("移动端", "390×844与430×932视口无横向溢出，操作区无遮挡"),
    ], [3.1, 13.5])
    doc.save(FILES["architecture"])


def build_review():
    doc = Document()
    style_document(doc, "BYU MANDARIN · 袈裟主题 · 教案审议")
    cover(doc, "PEDAGOGICAL REVIEW", "袈裟主题教学价值", "吸引力、文化深度与跨文化表达", "一件可观察的衣服承载修行身份、文学珍宝、收藏欲望和行动后果，适合中高级学习者形成有证据的口语解释。")
    doc.add_heading("一、教学价值", level=1)
    table(doc, ["维度", "价值", "课堂表现"], [
        ("具体", "材料、颜色、拼接、佩用者都可观察", "学生从可见特征开始说"),
        ("矛盾", "少欲提醒与珍贵宝物集中在同一件衣服", "学生使用本来、后来、可是解释"),
        ("角色", "金池长老与黑熊精体现不同占有理由", "学生追问、反驳、协商"),
        ("跨文化", "衣服常同时承担功能、身份与价格", "学生比较熟悉物件并说明边界"),
        ("语言", "常用词可以表达复杂判断", "认知深度高于词汇负担"),
    ], [3.0, 7.2, 6.4])
    doc.add_heading("二、文化深度", level=1)
    table(doc, ["3P", "内容", "核心问题"], [
        ("Product", "粪扫衣传统、拼接外观、文学中的锦襕袈裟", "同类衣物为什么呈现不同价值？"),
        ("Practice", "穿着、收藏、展示、借看、占有", "人怎样通过行动赋予物意义？"),
        ("Perspective", "少欲、身份、他人眼光、价值与欲望", "欲望来自物，还是来自看物的人？"),
    ], [3.0, 7.6, 6.0])
    doc.add_heading("三、跨文化比较", level=1)
    table(doc, ["物件", "可比较处", "差异"], [
        ("牛仔裤", "工作服进入大众时尚，也可代表身份与价格", "原始功能是耐穿；宗教意义不同"),
        ("限量运动鞋", "功能物品进入收藏与争抢", "商业稀缺与品牌机制更突出"),
        ("棒球帽", "遮阳用品可表示球队、地区和群体身份", "公共群体认同更突出"),
        ("学生自选", "连接个人文化经验", "需要说明选择依据和类比边界"),
    ], [3.0, 7.2, 6.4])
    doc.add_heading("四、预期产出", level=1)
    callout(doc, "达标", "学生说出袈裟的原来意义、人物后来赋予的价值、一条行动后果，以及一个跨文化物件的相似点和差异。")
    callout(doc, "发展充分", "学生使用两条证据，回应角色观点，指出类比边界，并在补说中调整中文表达。", "E2E7DD")
    doc.add_heading("五、研究观察", level=1)
    bullets(doc, [
        "识别文化物件向解释文化关系的发展。",
        "角色追问引发的口语长度、追问方式和自我修正。",
        "Hint使用前后话语在证据、因果和比较上的变化。",
        "即时解释与扩展阅读后解释之间的变化。",
        "学生对游戏叙事、文学文本与文化事实来源的区分。",
    ])
    doc.save(FILES["review"])


SOURCES = {
    "journey": "https://zh.wikisource.org/wiki/%E8%A5%BF%E9%81%8A%E8%A8%98/%E7%AC%AC016%E5%9B%9E",
    "museum": "https://www.kyohaku.go.jp/jp/collection/meihin/senshoku/item05/",
    "game": "https://www.jiemian.com/article/11631724.html",
    "jeans": "https://www.smithsonianmag.com/smithsonian-institution/the-origin-of-blue-jeans-89612175/",
    "denim": "https://www.smithsonianmag.com/arts-culture/denim-political-symbol-1960s-180976241/",
    "ap": "https://apnews.com/article/0cc01b7efb06bf27fd8c0c0cb2d1cc23",
}


def source(doc, title, kind, purpose, url):
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    shade(cell, PAPER)
    borders(cell)
    p = cell.paragraphs[0]
    run_style(p.add_run(title + "\n"), 10, True, CINNABAR)
    run_style(p.add_run(kind + "｜" + purpose + "\n"), 9, color=GOLD_DARK)
    run_style(p.add_run(url), 8.2, color="3F6485")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def build_teacher_readings():
    doc = Document()
    style_document(doc, "BYU MANDARIN · 袈裟主题 · 扩展阅读教师版")
    cover(doc, "EXTENSION READING · TEACHER", "袈裟主题扩展阅读材料包", "任务完成后持续开放的四条阅读路径", "学生每次选择一条材料，把新证据转化为60–90秒中文口语说明。")
    doc.add_heading("一、阅读结构", level=1)
    table(doc, ["路径", "材料", "问题", "建议"], [
        ("原著", "《西游记》第十六回", "喜欢在哪一个行动中转成伤害？", "IH–AL；5–8分钟"),
        ("文物", "博物馆袈裟藏品页", "外观、材料和意义怎样连接？", "IH–AL；4–6分钟"),
        ("改编", "游戏动画评论", "视觉改编怎样呈现欲望？", "AL；5–8分钟"),
        ("比较", "牛仔裤文化史与新闻", "功能怎样转成身份与价格？", "IH–AL；6–10分钟"),
    ], [2.5, 5.5, 6.3, 3.3])
    doc.add_paragraph("输入可使用中文、英文或图像；最终产出使用中文。学生在任务完成页、个人历史记录和教师指定的课后入口中持续访问材料。")

    doc.add_heading("二、四条路径", level=1)
    table(doc, ["路径", "教师导读", "口语任务", "词语"], [
        ("A 原著", "金池长老从借看走向藏留与纵火；黑熊精在火中取走袈裟。", "用本来、后来、因此说明行动变化。", "借看、归还、占有、纵火、后果"),
        ("B 文物", "一件藏品用精细织造模仿弃布拼接和针脚。", "说明看起来与实际上，并提出一个文化问题。", "弃布、缝成、模仿、织物、外观"),
        ("C 改编", "游戏动画用袈裟的移动、包围和火焰呈现金池长老的欲望。", "比较原著行动与游戏画面。", "改编、视觉、象征、保留、增加"),
        ("D 比较", "牛仔裤从劳动服进入时尚、身份与价格表达。", "说明一个相似点、一个差异和类比边界。", "耐穿、劳动者、身份、价格、类比"),
    ], [2.5, 6.7, 5.8, 3.1])

    doc.add_heading("三、口语提交", level=1)
    callout(doc, "统一提示", "我读到／看到……。这让我重新理解……。它和……有一点像，因为……。这个比较的边界在……。")
    table(doc, ["维度", "达标", "发展充分"], [
        ("证据", "准确使用一条材料", "连接两条来源"),
        ("解释", "说明人物行动或物件意义", "说明物、行动与价值观的关系"),
        ("比较", "相似点与差异", "类比帮助与类比边界"),
        ("中文", "形成连续话段", "能够补充、自我修正和追问"),
    ], [3.0, 6.8, 6.8])

    doc.add_heading("四、来源", level=1)
    source(doc, "《西游记》第十六回", "文学原典", "情节与人物行动", SOURCES["journey"])
    source(doc, "京都国立博物馆：犍陀穀糸袈裟", "博物馆藏品页", "粪扫衣传统与仿制外观", SOURCES["museum"])
    source(doc, "《黑神话：悟空》动画短片评论", "媒体评论", "游戏视觉改编", SOURCES["game"])
    source(doc, "Smithsonian: The Origin of Blue Jeans", "博物馆杂志", "工作服进入日常时尚", SOURCES["jeans"])
    source(doc, "Smithsonian: How Denim Became a Political Symbol of the 1960s", "历史文化文章", "阶级、性别与社会运动", SOURCES["denim"])
    source(doc, "AP 2026: Blue jeans have been a common thread", "通讯社新闻", "当代文化解释与多元身份", SOURCES["ap"])

    doc.add_heading("五、更新", level=1)
    bullets(doc, [
        "新闻链接每学期检查；历史与博物馆资料每学年检查。",
        "教师改写保留来源标题、机构、日期、类型和链接。",
        "学生页每条路径保留一项主材料，控制信息密度。",
        "扩展材料持续回到袈裟、价值、欲望与行动后果。",
    ])
    doc.save(FILES["teacher"])


def build_student_readings():
    doc = Document()
    style_document(doc, "BYU MANDARIN · 山外见闻 · 学生选读")
    cover(doc, "EXTENSION READING · STUDENT", "山外见闻：袈裟主题选读", "四选一 · 任务完成后随时回来", "选择一条材料，带着一条新证据回到故事，用中文说明你的新理解。")
    table(doc, ["选择", "你会发现", "读后要说", "时间"], [
        ("A 原著", "喜欢怎样变成占有和伤害", "哪个行动改变了事情？", "5–8分钟"),
        ("B 文物", "“像弃布”怎样成为一种外观", "材料和意义怎样连接？", "4–6分钟"),
        ("C 改编", "游戏怎样让欲望变成画面", "它保留和增加了什么？", "5–8分钟"),
        ("D 比较", "工作服怎样获得身份与价格", "哪里相似，哪里不同？", "6–10分钟"),
    ], [2.4, 5.5, 6.5, 3.1])
    doc.add_heading("A｜读原著", level=1)
    doc.add_paragraph("金池长老见到锦襕袈裟后，先要求借看，后来想把它留下，并接受纵火夺衣的主意。黑熊精看到袈裟后也把它带走。找出一个让喜欢转成伤害的行动。")
    source(doc, "《西游记》第十六回", "文学原典", "借看、归还、占有、纵火、后果", SOURCES["journey"])
    doc.add_heading("B｜看文物", level=1)
    doc.add_paragraph("博物馆的一件袈裟看起来像由许多旧布缝成。藏品用精细织造技术模仿拼布和针脚。思考人们为什么认真做出“像弃布”的样子。")
    source(doc, "京都国立博物馆袈裟藏品页", "博物馆物证", "弃布、缝成、模仿、织物、外观", SOURCES["museum"])
    doc.add_heading("C｜看游戏改编", level=1)
    doc.add_paragraph("评论者关注游戏动画中会移动、包围人物并与火焰连在一起的袈裟。比较原著中的人物行动和游戏增加的视觉画面。")
    source(doc, "游戏动画短片评论", "媒体评论", "改编、视觉、象征、保留、增加", SOURCES["game"])
    doc.add_heading("D｜比较文化物件", level=1)
    doc.add_paragraph("牛仔裤最初重视结实耐穿，后来也能表达时尚、反叛、身份或价格。你也可以选择限量运动鞋、棒球帽或自己文化中的物件。说明它和袈裟的相似点、差异和比较边界。")
    source(doc, "Smithsonian: The Origin of Blue Jeans", "博物馆杂志", "耐穿、劳动者、身份、价格、类比", SOURCES["jeans"])
    doc.add_heading("录一段60–90秒中文", level=1)
    callout(doc, "提示", "我读到／看到……。这让我重新理解……。它和……有一点像，因为……。这个比较的边界在……。")
    doc.add_paragraph("提交一条材料证据和你的解释。比较路径再加入一个相似点和一个差异。")
    doc.save(FILES["student"])


if __name__ == "__main__":
    DOCS.mkdir(parents=True, exist_ok=True)
    build_dialogue()
    build_lesson()
    build_architecture()
    build_review()
    build_teacher_readings()
    build_student_readings()
    for path in FILES.values():
        print(path)
