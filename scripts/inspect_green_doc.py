from pathlib import Path
from docx import Document

path = Path(r"D:\BYU中文文化任务网站_第二稿\output\doc\5000美元与5万棵树_GAI课程网站_专业审核版.docx")
doc = Document(path)
for i, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.strip()
    if text:
        print(f"P{i:03d}\t{paragraph.style.name}\t{text}")
for ti, table in enumerate(doc.tables):
    print(f"TABLE {ti} {len(table.rows)}x{len(table.columns)}")
    for ri, row in enumerate(table.rows):
        values = [" ".join(cell.text.split()) for cell in row.cells]
        print(f"T{ti}R{ri}\t" + " || ".join(values))
